import os
try:
    import docx
except ImportError:
    docx = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
from pathlib import Path
from typing import Dict, List, Any

from src.processors.base_processor import BaseProcessor
from src.utils.output_manager import OutputManager

# Import LLM functionality
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

import google.generativeai as genai
from data_extraction_pipeline import configure_gemini, retry_with_backoff

class AdditionalDocProcessor(BaseProcessor):
    """Processes additional documents (transcripts, emails, updates)"""
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.doc', '.docx', '.txt']
        # Configure Gemini API
        try:
            configure_gemini()
        except Exception as e:
            print(f"Warning: Could not configure Gemini API: {e}")
    
    def get_supported_extensions(self) -> List[str]:
        """Return list of supported file extensions"""
        return self.supported_extensions
    
    def process(self, file_path: str, output_dir: str) -> Dict[str, Any]:
        """
        Process additional documents with simplified pipeline
        
        Args:
            file_path: Path to the document file
            output_dir: Directory to save processed outputs
            
        Returns:
            Dictionary containing processing results
        """
        try:
            print(f"Processing additional document: {file_path}")
            
            # Extract text content
            text_content = self._extract_text(file_path)
            
            if not text_content or not text_content.strip():
                raise ValueError("Could not extract text content from the document")
            
            # Structure content using LLM
            filename = Path(file_path).name
            structured_content = self._structure_content(text_content, filename)
            
            # Convert to markdown
            markdown_content = self._text_to_markdown(structured_content, filename)
            
            # Get output path
            output_paths = OutputManager.get_output_paths(output_dir, 'additional', filename)
            
            # Save markdown file
            success = OutputManager.save_file(markdown_content, output_paths['markdown'])
            
            if success:
                return {
                    'status': 'success',
                    'filename': filename,
                    'output_file': output_paths['markdown'],
                    'content_length': len(text_content)
                }
            else:
                raise ValueError("Failed to save processed content")
                
        except Exception as e:
            print(f"Error processing additional document: {e}")
            return {
                'status': 'error',
                'filename': Path(file_path).name,
                'error': str(e)
            }
    
    def _extract_text(self, file_path: str) -> str:
        """Extract raw text based on file type"""
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.txt':
                return self._extract_text_from_txt(file_path)
            elif file_extension == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                return self._extract_text_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            print(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            import fitz  # PyMuPDF (already available)
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            
            doc.close()
            return text
            
        except Exception as e:
            print(f"Error extracting text from PDF using PyMuPDF: {e}")
            # Fallback to PyPDF2 if available
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()
                    return text
            except Exception as e2:
                print(f"Error with PyPDF2 fallback: {e2}")
                return ""
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not docx:
            raise ImportError("python-docx not available")
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
            
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""
    
    @retry_with_backoff()
    def _structure_content(self, text: str, filename: str) -> str:
        """Use LLM to structure and clean up content"""
        try:
            model = genai.GenerativeModel(os.getenv("GEMINI_MODEL", "gemini-1.5-flash"))
            
            prompt = f"""
            Please analyze and structure the following document content. The document is named "{filename}".
            
            Your task:
            1. Clean up and organize the content
            2. Identify the document type (e.g., call transcript, email, update, report)
            3. Extract key information and organize it logically
            4. Maintain important details while improving readability
            5. If it's a transcript, identify speakers and structure conversations
            6. If it's an email, preserve sender/receiver and subject information
            7. If it's an update or report, organize by topics/sections
            
            Document content:
            {text[:8000]}  # Limit to first 8000 characters to avoid token limits
            
            Please provide a well-structured, cleaned version of this content.
            """
            
            response = model.generate_content(prompt)
            return response.text
            
        except Exception as e:
            print(f"Error structuring content with LLM: {e}")
            # Return original text if LLM processing fails
            return text
    
    def _text_to_markdown(self, structured_text: str, filename: str) -> str:
        """Convert structured text to markdown format"""
        markdown_content = []
        
        # Add document header
        doc_name = Path(filename).stem.replace('_', ' ').replace('-', ' ').title()
        markdown_content.append(f"# {doc_name}")
        markdown_content.append("")
        
        # Add metadata
        markdown_content.append("## Document Information")
        markdown_content.append("")
        markdown_content.append(f"**Original Filename:** {filename}")
        markdown_content.append(f"**Document Type:** Additional Document")
        markdown_content.append("")
        
        # Add structured content
        markdown_content.append("## Content")
        markdown_content.append("")
        
        # If the structured text doesn't already have markdown formatting, add basic structure
        if not any(line.startswith('#') for line in structured_text.split('\n')):
            # Add the content as-is but ensure proper formatting
            lines = structured_text.split('\n')
            for line in lines:
                if line.strip():
                    markdown_content.append(line)
                else:
                    markdown_content.append("")
        else:
            # Content already has markdown formatting
            markdown_content.append(structured_text)
        
        return "\n".join(markdown_content)