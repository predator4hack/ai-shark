"""
Document Converter Utility for AI-Shark

Converts markdown files to DOCX format for download functionality.
"""

import os
import re
from pathlib import Path
from typing import Optional, Dict, Any
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None


class MarkdownToDocxConverter:
    """
    Converter class for transforming markdown files to DOCX format
    """
    
    def __init__(self):
        """Initialize the converter"""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx package is required for DOCX conversion. "
                "Install it with: pip install python-docx"
            )
        
        self.converter_name = "MarkdownToDocxConverter"
    
    def convert_markdown_file(self, markdown_file: str, docx_file: Optional[str] = None) -> str:
        """
        Convert a markdown file to DOCX format
        
        Args:
            markdown_file: Path to the markdown file
            docx_file: Optional output path for DOCX file. If None, uses same name with .docx extension
            
        Returns:
            Path to the generated DOCX file
            
        Raises:
            FileNotFoundError: If markdown file doesn't exist
            Exception: If conversion fails
        """
        markdown_path = Path(markdown_file)
        
        if not markdown_path.exists():
            raise FileNotFoundError(f"Markdown file not found: {markdown_file}")
        
        # Determine output file path
        if docx_file is None:
            docx_file = str(markdown_path.with_suffix('.docx'))
        
        print(f"🔄 Converting markdown to DOCX")
        print(f"   📝 Input: {markdown_file}")
        print(f"   📄 Output: {docx_file}")
        
        try:
            # Read markdown content
            with open(markdown_file, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            # Convert to DOCX
            self.convert_markdown_content(markdown_content, docx_file)
            
            print(f"✅ DOCX conversion completed: {docx_file}")
            return docx_file
            
        except Exception as e:
            print(f"❌ DOCX conversion failed: {e}")
            raise
    
    def convert_markdown_content(self, markdown_content: str, docx_file: str) -> None:
        """
        Convert markdown content to DOCX document
        
        Args:
            markdown_content: The markdown content to convert
            docx_file: Path for the output DOCX file
        """
        # Create new document
        doc = Document()
        
        # Configure document styles
        self._setup_document_styles(doc)
        
        # Parse and convert markdown content
        lines = markdown_content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                # Empty line - add space
                doc.add_paragraph()
                i += 1
                continue
            
            # Process different markdown elements
            if line.startswith('# '):
                # H1 - Main heading
                self._add_heading(doc, line[2:], level=1)
            elif line.startswith('## '):
                # H2 - Section heading
                self._add_heading(doc, line[3:], level=2)
            elif line.startswith('### '):
                # H3 - Subsection heading
                self._add_heading(doc, line[4:], level=3)
            elif line.startswith('#### '):
                # H4 - Minor heading
                self._add_heading(doc, line[5:], level=4)
            elif line.startswith('**') and line.endswith('**'):
                # Bold line
                self._add_bold_paragraph(doc, line[2:-2])
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet list
                i = self._add_bullet_list(doc, lines, i)
                continue
            elif re.match(r'^\d+\.', line):
                # Numbered list
                i = self._add_numbered_list(doc, lines, i)
                continue
            elif line.startswith('---'):
                # Horizontal rule - add page break or spacing
                doc.add_page_break()
            else:
                # Regular paragraph
                self._add_paragraph(doc, line)
            
            i += 1
        
        # Save document
        doc.save(docx_file)
    
    def _setup_document_styles(self, doc: Document) -> None:
        """Setup custom styles for the document"""
        try:
            # Create custom styles if they don't exist
            styles = doc.styles
            
            # Header style
            if 'AI-Shark Header' not in [s.name for s in styles]:
                header_style = styles.add_style('AI-Shark Header', WD_STYLE_TYPE.PARAGRAPH)
                header_style.font.name = 'Calibri'
                header_style.font.size = Pt(14)
                header_style.font.bold = True
                header_style.paragraph_format.space_after = Pt(12)
            
            # Question style
            if 'AI-Shark Question' not in [s.name for s in styles]:
                question_style = styles.add_style('AI-Shark Question', WD_STYLE_TYPE.PARAGRAPH)
                question_style.font.name = 'Calibri'
                question_style.font.size = Pt(11)
                question_style.font.bold = True
                question_style.paragraph_format.left_indent = Inches(0.25)
                question_style.paragraph_format.space_before = Pt(6)
                question_style.paragraph_format.space_after = Pt(6)
            
        except Exception as e:
            print(f"⚠️ Warning: Could not create custom styles: {e}")
    
    def _add_heading(self, doc: Document, text: str, level: int = 1) -> None:
        """Add a heading to the document"""
        heading = doc.add_heading(text, level=level)
        
        # Customize heading formatting
        if level == 1:
            heading.paragraph_format.space_before = Pt(18)
            heading.paragraph_format.space_after = Pt(12)
        elif level == 2:
            heading.paragraph_format.space_before = Pt(14)
            heading.paragraph_format.space_after = Pt(8)
        else:
            heading.paragraph_format.space_before = Pt(10)
            heading.paragraph_format.space_after = Pt(6)
    
    def _add_paragraph(self, doc: Document, text: str) -> None:
        """Add a regular paragraph to the document"""
        # Clean up markdown formatting
        text = self._clean_markdown_formatting(text)
        
        paragraph = doc.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.15
    
    def _add_bold_paragraph(self, doc: Document, text: str) -> None:
        """Add a bold paragraph to the document"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        paragraph.paragraph_format.space_after = Pt(6)
    
    def _add_bullet_list(self, doc: Document, lines: list, start_index: int) -> int:
        """Add a bullet list to the document"""
        i = start_index
        
        while i < len(lines):
            line = lines[i].strip()
            
            if line.startswith('- ') or line.startswith('* '):
                # Add bullet point
                text = line[2:].strip()
                text = self._clean_markdown_formatting(text)
                
                paragraph = doc.add_paragraph(text, style='List Bullet')
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.space_after = Pt(3)
                
                i += 1
            else:
                # End of list
                break
        
        return i
    
    def _add_numbered_list(self, doc: Document, lines: list, start_index: int) -> int:
        """Add a numbered list to the document"""
        i = start_index
        
        while i < len(lines):
            line = lines[i].strip()
            
            if re.match(r'^\d+\.', line):
                # Add numbered point
                text = re.sub(r'^\d+\.\s*', '', line)
                text = self._clean_markdown_formatting(text)
                
                paragraph = doc.add_paragraph(text, style='List Number')
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.space_after = Pt(3)
                
                i += 1
            else:
                # End of list
                break
        
        return i
    
    def _clean_markdown_formatting(self, text: str) -> str:
        """Clean markdown formatting from text"""
        # Remove bold/italic markers
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)      # Italic
        text = re.sub(r'`(.*?)`', r'\1', text)        # Inline code
        
        # Clean up links - keep just the text
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        
        return text


def convert_founders_checklist_to_docx(markdown_file: str, 
                                     docx_file: Optional[str] = None) -> str:
    """
    Convert a founders checklist markdown file to DOCX format
    
    Args:
        markdown_file: Path to the markdown file
        docx_file: Optional output path for DOCX file
        
    Returns:
        Path to the generated DOCX file
        
    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If markdown file doesn't exist
        Exception: If conversion fails
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx package is required for DOCX conversion. "
            "Install it with: pip install python-docx"
        )
    
    converter = MarkdownToDocxConverter()
    return converter.convert_markdown_file(markdown_file, docx_file)


def is_docx_conversion_available() -> bool:
    """
    Check if DOCX conversion is available
    
    Returns:
        True if python-docx is installed and conversion is available
    """
    return DOCX_AVAILABLE


def get_docx_converter_info() -> Dict[str, Any]:
    """
    Get information about the DOCX converter
    
    Returns:
        Dictionary with converter information
    """
    return {
        "available": DOCX_AVAILABLE,
        "required_package": "python-docx",
        "install_command": "pip install python-docx",
        "supported_formats": ["markdown", "md"],
        "output_format": "docx"
    }